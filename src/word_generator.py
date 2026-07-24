from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_word_document(pump_data):

    document = Document()

    # ===========================
    # Main Title
    # ===========================

    title = document.add_heading("TECHNICAL OFFER", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_paragraph()

    # ===========================
    # Add each section
    # ===========================

    for section, fields in pump_data.items():

        heading = document.add_heading(section.upper(), level=2)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        for parameter, value in fields.items():

            if value is None:
                continue

            row = table.add_row().cells

            row[0].text = str(parameter)
            row[1].text = str(value)

            # Make parameter text bold
            for run in row[0].paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10)

            for run in row[1].paragraphs[0].runs:
                run.font.size = Pt(10)

        document.add_paragraph()

    document.save("output/technical_offer.docx")